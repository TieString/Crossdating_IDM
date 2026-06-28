"""Build the editable Word version of docs/technical-note-manuscript.md.

This intentionally uses a small Markdown subset so that the manuscript remains
reviewable as text in the repository while the DOCX is regenerated consistently.
"""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "technical-note-manuscript.md"
OUTPUT = ROOT / "docs" / "Crossdating_IDM_Technical_Note.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, size: float, color=BLACK, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style(style, font_size, color, before, after, line_spacing, bold=False, alignment=None):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(font_size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing
    if alignment is not None:
        style.paragraph_format.alignment = alignment


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), "B7C9DD")
    borders.append(bottom)
    p_pr.append(borders)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    """Apply fixed 9360-DXA geometry and the selected preset's cell margins."""
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def add_markdown_table(doc, lines):
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    column_count = len(rows[0])
    rows = [row[:column_count] + [""] * max(0, column_count - len(row)) for row in rows]
    widths = [2100, 3800, 3460] if column_count == 3 else [9360 // column_count] * column_count
    if sum(widths) != 9360:
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(value)
            set_run_font(r, 9.3, BLACK, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, "F2F4F7")
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style(normal, 11, BLACK, 0, 8, 1.333, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    set_style(doc.styles["Heading 1"], 16, BLUE, 18, 10, 1.0, bold=True)
    set_style(doc.styles["Heading 2"], 13, BLUE, 12, 6, 1.0, bold=True)
    set_style(doc.styles["Heading 3"], 12, DARK_BLUE, 8, 4, 1.0, bold=True)

    metadata = doc.styles.add_style("Technical Note Metadata", WD_STYLE_TYPE.PARAGRAPH)
    set_style(metadata, 10, MUTED, 0, 2, 1.1)

    abstract = doc.styles.add_style("Technical Note Abstract", WD_STYLE_TYPE.PARAGRAPH)
    set_style(abstract, 10.5, BLACK, 0, 8, 1.15)
    abstract.paragraph_format.left_indent = Inches(0.15)
    abstract.paragraph_format.right_indent = Inches(0.15)

    references = doc.styles.add_style("Technical Note References", WD_STYLE_TYPE.PARAGRAPH)
    set_style(references, 10, BLACK, 0, 5, 1.15)
    references.paragraph_format.left_indent = Inches(0.25)
    references.paragraph_format.first_line_indent = Inches(-0.25)

    code = doc.styles.add_style("Technical Note Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(9.5)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)

    caption = doc.styles["Caption"]
    set_style(caption, 9.5, MUTED, 4, 10, 1.1)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("TECHNICAL NOTE | CROSSDATING IDM")
    set_run_font(run, 8.5, MUTED, bold=True)
    add_bottom_border(p)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Draft for author completion | Page ")
    set_run_font(run, 8.5, MUTED)
    add_page_field(p)


def add_title_block(doc, title, metadata):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(4)
    kicker.paragraph_format.space_after = Pt(7)
    r = kicker.add_run("TECHNICAL NOTE")
    set_run_font(r, 10.5, DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_run_font(r, 22, BLACK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("Software for auditable RWL editing, COFECHA execution, and constrained diagnostic review")
    set_run_font(r, 12, MUTED, italic=True)

    for line in metadata:
        p = doc.add_paragraph(style="Technical Note Metadata")
        label, value = line.split(":", 1)
        r = p.add_run(f"{label}: ")
        set_run_font(r, 10, MUTED, bold=True)
        r = p.add_run(value.strip())
        set_run_font(r, 10, MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(7)
    rule.paragraph_format.space_after = Pt(11)
    add_bottom_border(rule)


def split_markdown_label(line):
    """Split a line written as **Label:** value without leaking Markdown marks."""
    label_end = line.find(":**")
    return line[2:label_end], line[label_end + 3:].strip()


def add_text_paragraph(doc, text, style="Normal"):
    paragraph = doc.add_paragraph(style=style)
    # Preserve only the light emphasis needed by the manuscript source.
    cursor = 0
    while cursor < len(text):
        start = text.find("`", cursor)
        if start == -1:
            run = paragraph.add_run(text[cursor:])
            set_run_font(run, 9.5 if style == "Technical Note Code" else (10 if style == "Technical Note References" else 11))
            break
        if start > cursor:
            run = paragraph.add_run(text[cursor:start])
            set_run_font(run, 11)
        end = text.find("`", start + 1)
        if end == -1:
            run = paragraph.add_run(text[start:])
            set_run_font(run, 11)
            break
        run = paragraph.add_run(text[start + 1:end])
        set_run_font(run, 10, DARK_BLUE)
        cursor = end + 1
    return paragraph


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    title = ""
    metadata = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("# "):
            title = line[2:]
        elif line.startswith("**") and ":**" in line:
            label, value = split_markdown_label(line)
            metadata.append(f"{label}: {value}")
        elif line.startswith("## Abstract"):
            i += 1
            break
        i += 1

    add_title_block(doc, title, metadata)
    abstract_heading = doc.add_paragraph(style="Heading 2")
    abstract_heading.add_run("Abstract")
    current_section = ""
    in_references = False
    in_code = False

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        i += 1
        if not line:
            continue
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            add_text_paragraph(doc, raw, "Technical Note Code")
            continue
        if line.startswith("## "):
            current_section = line[3:]
            in_references = current_section == "References"
            heading = doc.add_paragraph(style="Heading 1")
            heading.add_run(current_section)
            continue
        if line.startswith("### "):
            heading = doc.add_paragraph(style="Heading 2")
            heading.add_run(line[4:])
            continue
        image_match = re.fullmatch(r"!\[[^\]]*\]\(([^)]+)\)", line)
        if image_match:
            image_path = SOURCE.parent / image_match.group(1)
            doc.add_picture(str(image_path), width=Inches(6.22))
            continue
        if line.startswith("|") and line.endswith("|"):
            table_lines = [line]
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        if line.startswith("Fig. "):
            add_text_paragraph(doc, line, "Caption")
            continue
        if line.startswith("**Keywords:**"):
            p = doc.add_paragraph(style="Technical Note Abstract")
            label = p.add_run("Keywords: ")
            set_run_font(label, 10.5, BLACK, bold=True)
            rest = p.add_run(line[len("**Keywords:**"):].strip())
            set_run_font(rest, 10.5, BLACK)
            continue
        if line.startswith("**") and ":**" in line:
            p = doc.add_paragraph(style="Technical Note Metadata")
            label, value = split_markdown_label(line)
            r = p.add_run(f"{label}: ")
            set_run_font(r, 10, MUTED, bold=True)
            r = p.add_run(value.strip())
            set_run_font(r, 10, MUTED)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            r = p.add_run(line[2:])
            set_run_font(r, 11)
            continue
        add_text_paragraph(
            doc,
            line,
            "Technical Note References" if in_references else "Technical Note Abstract" if current_section == "Abstract" else "Normal",
        )

    doc.core_properties.title = title
    doc.core_properties.subject = "Technical Note"
    doc.core_properties.author = "To be supplied by submitting authors"
    doc.core_properties.comments = "Generated from docs/technical-note-manuscript.md"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
